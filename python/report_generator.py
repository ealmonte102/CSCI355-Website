import os
import sys
import docx
import json
import re;
from docx.shared import Inches

TEMPLATE_PATH = '../assets/hard_submission_template.docx'
PATH_DIRECTORY = os.getcwd()

class ReportGenerator():
    @staticmethod
    def generateReport(config):
        document = docx.Document(TEMPLATE_PATH)
        document.add_paragraph('CSCI 323-33')
        programParagraph = document.add_paragraph(
            f'Program: Project #{config["project_num"]} {config["project_title"]} ')
        programParagraph.add_run('<{}>'.format(config["language"])).bold = True
        document.add_paragraph(f'Name: {config["name"]}')
        softCopyDateParagraph = document.add_paragraph(f'Soft copy: {config["soft_copy_due_date"]}')
        softCopyDateParagraph.paragraph_format.left_indent = Inches(0.25)
        hardCopyDateParagraph = document.add_paragraph(f'Hard copy: {config["hard_copy_due_date"]}')
        hardCopyDateParagraph.paragraph_format.left_indent = Inches(0.25)
        document.add_paragraph("Algorithm Steps").style = document.styles['CSCI 323 Section Heading']

        for step in config["algorithm_steps"].splitlines():
            underline = False
            bold = False
            pattern = re.compile("^\\s*\\d*\\.?\\d+\\)")
            run = None
            if step.startswith('__'):
                step = step[2:step.find('__', 2)]
                underline = True
            if step.startswith('#'):
                step = step[1:step.find('#', 1)]
                bold = True;
            if pattern.match(step):
                bullet = step[:step.find(")")]
                step = step[step.find(")") + 1:].strip()
                if step.startswith('__'):
                    step = step[2:step.find('__', 2)]
                    underline = True
                if step.startswith('#'):
                    step = step[1:step.find('#', 1)]
                    bold = True;
                p = document.add_paragraph()
                p.add_run(bullet)
                p.add_run(") ")
                run = p.add_run(step)
            else:
                p = document.add_paragraph(step)
                run = p.runs[0];
            if underline:
                run.underline = True
            if bold:
                run.bold = True
        document.add_page_break()
        document.add_paragraph('Input', 'CSCI 323 Section Heading')
        for inputFile in config["input_files"]:
            header = document.add_paragraph(inputFile["name"], 'CSCI 323 Section Heading 2')
            fileText = open(inputFile["path"], 'r').read()
            document.add_paragraph(fileText, header.style.next_paragraph_style)
        document.add_page_break()
        document.add_paragraph('Output', 'CSCI 323 Section Heading')
        for outputFile in config["output_files"]:
            header = document.add_paragraph(outputFile["name"],
                                            'CSCI 323 Section Heading 2')
            fileText = open(outputFile["path"], 'r').read()
            try:
                document.add_paragraph(fileText, header.style.next_paragraph_style)
            except ValueError as valueError:
                print("Error reading file: " + outputFile["name"])
                print(str(valueError))
        document.add_page_break()
        document.add_paragraph('Source Code', 'CSCI 323 Section Heading')
        currentSourceCodeStyle = document.styles['Source Code 2']
        sourceCodeFile = config["source_code_file"]
        document.add_paragraph(sourceCodeFile["name"],
                               'CSCI 323 Section Heading 2')
        for line in open(sourceCodeFile["path"], 'r').read().splitlines():
            document.add_paragraph(line, currentSourceCodeStyle.next_paragraph_style)
            currentSourceCodeStyle = currentSourceCodeStyle.next_paragraph_style
        documentFilename = f'{config["output"]}.docx'
        document.save(documentFilename)
        return documentFilename


if __name__ == "__main__":
    reportGenerator = ReportGenerator()
    config = json.load(open(sys.argv[1]))
    sys.stdout.write(reportGenerator.generateReport(config))